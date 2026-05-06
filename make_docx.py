from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.15)
    section.right_margin = Inches(1.15)

styles = doc.styles

styles['Title'].font.name = 'Calibri'
styles['Title'].font.size = Pt(22)
styles['Title'].font.bold = True
styles['Title'].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

styles['Heading 2'].font.name = 'Calibri'
styles['Heading 2'].font.size = Pt(12)
styles['Heading 2'].font.bold = True
styles['Heading 2'].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

styles['Heading 3'].font.name = 'Calibri'
styles['Heading 3'].font.size = Pt(10.5)
styles['Heading 3'].font.bold = True
styles['Heading 3'].font.color.rgb = RGBColor(0x40, 0x40, 0x40)
styles['Heading 3'].font.italic = False

styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(10.5)

styles['List Bullet'].font.name = 'Calibri'
styles['List Bullet'].font.size = Pt(10.5)


def add_bullet(doc, text):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(2)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)
    return para


# ---- Parse study_guide.md ----
with open(r'c:\Users\pamurphy\Workspace\literacy_month\study_guide.md', encoding='utf-8') as f:
    lines = f.readlines()

# Title
t = doc.add_paragraph('SIERRA Literacy Month \u2014 Study Guide', style='Title')
t.alignment = WD_ALIGN_PARAGRAPH.LEFT

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(14)
run = sub.add_run('Certification Quiz Prep  |  22 Days')
run.italic = True
run.font.size = Pt(10.5)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    if line.startswith('# '):
        i += 1
        continue

    if line.startswith('## '):
        heading = line[3:].strip()
        doc.add_heading(heading, level=2)
        i += 1
        continue

    if line.startswith('**Know these ideas**') or line.startswith('**Work application**'):
        label = line.strip().strip('*')
        h = doc.add_heading(label, level=3)
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(2)
        i += 1
        continue

    if line.startswith('- '):
        bullet_text = line[2:].strip()
        add_bullet(doc, bullet_text)
        i += 1
        continue

    if line.startswith('---') or line.strip() == '':
        i += 1
        continue

    # Plain paragraph (misconceptions header etc.)
    if line.strip():
        para = doc.add_paragraph(style='Normal')
        para.paragraph_format.space_after = Pt(4)
        parts = re.split(r'(\*\*[^*]+\*\*)', line.strip())
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)

    i += 1

out = r'c:\Users\pamurphy\Workspace\literacy_month\SIERRA_Study_Guide.docx'
doc.save(out)
print('saved:', out)
